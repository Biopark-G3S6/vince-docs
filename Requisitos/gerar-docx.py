#!/usr/bin/env python3
"""Converte um documento Markdown deste repositório em .docx formatado para publicação.

O .docx é artefato derivado: a fonte da verdade é sempre o Markdown. Ver README.md,
seção "Documentos oficiais".

Dependência: python-docx. Em ambiente isolado, para não sujar o Python do sistema:

    python3 -m venv .venv
    .venv/bin/pip install python-docx

Uso:

    .venv/bin/python Requisitos/gerar-docx.py Requisitos/URS.md Requisitos/URS.docx

Formatado para importar limpo no Google Docs: fonte monoespaçada nativa do catálogo do
Google, sumário estático (o campo TOC do Word chega vazio na importação) e diagramas
dimensionados para não quebrar linha em A4.
"""

import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)
SUBHEADING_COLOR = RGBColor(0x2E, 0x5A, 0x88)
CODE_COLOR = RGBColor(0xC7, 0x25, 0x4E)
CODE_FILL = "F7F2F4"
TABLE_HEADER_FILL = "DEE6F1"
LABEL_FILL = "EDEFF3"
CODEBLOCK_FILL = "F5F5F5"

# Courier New existe nativamente no Google Docs; Consolas não, e seria substituída na importação.
MONO_FONT = "Courier New"

BODY_FONT = "Arial"

# O padrão institucional não usa hierarquia visual de cor; os quadros de requisito e as
# tabelas carregam a estrutura sozinhos.
CAPA_CAMPOS = ("Disciplina", "Professor", "Acadêmicos")
FICHA_CAMPOS = ("Projeto", "Cliente", "Versão", "Data", "Status")
FIELD_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.*)$")
RF_TITLE_RE = re.compile(r"^(RF-[A-Z]{3}-\d{3})\s+—\s+(.*)$")
RF_REF_RE = re.compile(r"RF-[A-Z]{3}-\d{3}")

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]*\))")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)-\s+(.*)$")
ORDERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")


# --------------------------------------------------------------------------- blocos


def parse_blocks(lines):
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i].rstrip("\n"))
                i += 1
            i += 1
            blocks.append(("code", buf))
            continue

        m = HEADING_RE.match(stripped)
        if m:
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and TABLE_SEP_RE.match(lines[i + 1].strip()):
            header = split_row(stripped)
            aligns = [cell_align(c) for c in split_row(lines[i + 1].strip())]
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            blocks.append(("table", (header, aligns, rows)))
            continue

        if BULLET_RE.match(line) or ORDERED_RE.match(line):
            items = []
            while i < n:
                cur = lines[i].rstrip("\n")
                if not cur.strip():
                    break
                mb = BULLET_RE.match(cur)
                mo = ORDERED_RE.match(cur)
                if mb:
                    items.append([len(mb.group(1)) // 2, False, mb.group(2).strip()])
                elif mo:
                    items.append([len(mo.group(1)) // 2, True, mo.group(3).strip()])
                elif cur.startswith("  ") and items:
                    items[-1][2] += " " + cur.strip()
                else:
                    break
                i += 1
            blocks.append(("list", items))
            continue

        buf = []
        while i < n:
            cur = lines[i].rstrip("\n")
            s = cur.strip()
            if not s or s == "---" or HEADING_RE.match(s) or s.startswith("|") or s.startswith("```"):
                break
            if BULLET_RE.match(cur) or ORDERED_RE.match(cur):
                break
            buf.append(s)
            i += 1
        if buf:
            blocks.append(("para", buf))
    return blocks


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def cell_align(sep):
    sep = sep.strip()
    if sep.startswith(":") and sep.endswith(":"):
        return WD_ALIGN_PARAGRAPH.CENTER
    if sep.endswith(":"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


# --------------------------------------------------------------------------- runs


def shade_run(run, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    run._element.get_or_add_rPr().append(shd)


def add_inline(paragraph, text, bold=False, size=None):
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            add_inline(paragraph, token[2:-2], bold=True, size=size)
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(9) if size is None else size - Pt(1)
            run.font.color.rgb = CODE_COLOR
            run.bold = bold
            shade_run(run, CODE_FILL)
        elif token.startswith("["):
            label = token[1 : token.index("]")]
            add_inline(paragraph, label, bold=bold, size=size)
        else:
            run = paragraph.add_run(token)
            run.bold = bold
            if size is not None:
                run.font.size = size


# --------------------------------------------------------------------------- estilos


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    spec = {
        "Title": (24, HEADING_COLOR, 0, 6),
        "Heading 1": (16, HEADING_COLOR, 20, 8),
        "Heading 2": (13, HEADING_COLOR, 14, 6),
        "Heading 3": (11.5, SUBHEADING_COLOR, 12, 4),
        "Heading 4": (10.5, SUBHEADING_COLOR, 10, 4),
    }
    for name, (size, color, before, after) in spec.items():
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = BODY_FONT
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(2)


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, end):
        run._r.append(el)
    return run


def add_footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("VinceArt · URS 0.1 · ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        field = add_field(p, "PAGE")
        field.font.size = Pt(8)
        field.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_toc(doc, headings):
    """Sumário estático: campo TOC do Word chega vazio no Google Docs e parece defeito."""
    doc.add_heading("Índice", level=1)
    for level, text in headings:
        if level > 3:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0 if level == 2 else 0.7)
        run = p.add_run(text)
        run.bold = level == 2
        run.font.size = Pt(10.5 if level == 2 else 10)
        if level > 2:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_page_break()


# ----------------------------------------------------------------- capa e quadros


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def parse_meta(linhas):
    """Bloco de metadados da capa. Linha continuada pertence ao campo anterior — a
    detecção não pode depender de todas as linhas começarem em negrito."""
    meta, atual = {}, None
    for linha in linhas:
        m = FIELD_RE.match(linha)
        if m:
            atual = m.group(1)
            meta[atual] = m.group(2)
        elif atual:
            meta[atual] += " " + linha.strip()
        else:
            return None
    return meta if len(meta) >= 3 else None


def render_capa(doc, meta, logo_path):
    """Capa no padrão institucional: marca, identificação da disciplina e ficha do projeto."""
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_path, width=Cm(7.32))

    for campo in CAPA_CAMPOS:
        if campo not in meta:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, f"**{campo}:** {meta[campo]}")

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(p, meta.get("__titulo__", "Especificação de Requisitos"))

    if "Projeto" in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, f"**{meta['Projeto']}**", size=Pt(14))

    for _ in range(4):
        doc.add_paragraph()

    ficha = [(c, meta[c]) for c in FICHA_CAMPOS if c in meta]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for rotulo, valor in ficha:
        cells = table.add_row().cells
        cells[0].width = Cm(4.5)
        cells[1].width = Cm(11.4)
        for idx, (texto, negrito) in enumerate(((f"{rotulo}:", True), (valor, False))):
            cell = cells[idx]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            add_inline(par, texto, bold=negrito)
        shade(cells[0], LABEL_FILL)

    doc.add_page_break()


def _fields(items):
    """Agrupa a lista de um requisito em (rótulo, linhas), na ordem em que aparecem."""
    campos = []
    for level, ordered, text in items:
        m = FIELD_RE.match(text) if level == 0 else None
        if m:
            campos.append([m.group(1), ([(0, False, m.group(2))] if m.group(2) else [])])
        elif campos:
            campos[-1][1].append((level, ordered, text))
    return campos


def _fill_cell(cell, linhas):
    contadores = {}
    primeiro = True
    for level, ordered, text in linhas:
        par = cell.paragraphs[0] if primeiro else cell.add_paragraph()
        primeiro = False
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.line_spacing = 1.0
        if level:
            par.paragraph_format.left_indent = Cm(0.45 * level)
            if ordered:
                contadores[level] = contadores.get(level, 0) + 1
                marcador = f"{contadores[level]}. "
            else:
                contadores.pop(level, None)
                marcador = "• " if level == 1 else "– "
            par.paragraph_format.first_line_indent = Cm(-0.45)
            text = marcador + text
        else:
            contadores.clear()
        add_inline(par, text, size=Pt(9.5))


def render_requisito(doc, rf_id, nome, items):
    """Um requisito por quadro, no formato da especificação institucional."""
    campos = _fields(items)
    valores = {rotulo: linhas for rotulo, linhas in campos}

    def texto_de(rotulo):
        return " ".join(t for _, _, t in valores.get(rotulo, []))

    relacionados = sorted(
        {r for rotulo, linhas in campos for _, _, t in linhas for r in RF_REF_RE.findall(t)}
        - {rf_id}
    )

    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cabecalho = table.add_row().cells
    for idx, (texto, largura, rotulo) in enumerate((
        ("ID:", Cm(2.6), True), (rf_id, Cm(5.0), False),
        ("Prioridade:", Cm(2.6), True), (texto_de("Prioridade") or "—", Cm(5.7), False),
    )):
        cabecalho[idx].width = largura
        cabecalho[idx].text = ""
        par = cabecalho[idx].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        add_inline(par, texto, bold=True, size=Pt(9.5))
        if rotulo:
            shade(cabecalho[idx], LABEL_FILL)

    def linha(rotulo, linhas):
        cells = table.add_row().cells
        cells[1].merge(cells[3])
        cells[0].width = Cm(2.6)
        cells[0].text = ""
        par = cells[0].paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        add_inline(par, rotulo, bold=True, size=Pt(9.5))
        shade(cells[0], LABEL_FILL)
        cells[1].text = ""
        _fill_cell(cells[1], linhas)

    linha("Requisito:", [(0, False, f"**{nome}**")])
    for rotulo, linhas in campos:
        if rotulo == "Prioridade":
            continue
        linha(f"{rotulo}:", linhas)
    linha("Requisitos relacionados:", [(0, False, ", ".join(relacionados) if relacionados else "N/A")])

    doc.add_paragraph()


# --------------------------------------------------------------------------- render


def render_table(doc, header, aligns, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = aligns[idx] if idx < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)
        add_inline(p, text, bold=True, size=Pt(9.5))
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), TABLE_HEADER_FILL)
        cell._tc.get_or_add_tcPr().append(shd)

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row[: len(header)]):
            cell = cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = aligns[idx] if idx < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text, size=Pt(9.5))
    doc.add_paragraph()


def render_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODEBLOCK_FILL)
    cell._tc.get_or_add_tcPr().append(shd)
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(7.5)
    doc.add_paragraph()


BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


def render_list(doc, items):
    for level, ordered, text in items:
        level = min(level, 2)
        styles = NUMBER_STYLES if ordered else BULLET_STYLES
        try:
            p = doc.add_paragraph(style=styles[level])
        except KeyError:
            p = doc.add_paragraph(style=styles[0])
            p.paragraph_format.left_indent = Cm(0.8 * (level + 1))
        add_inline(p, text)


def render(blocks, out_path, logo_path=None):
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    doc.core_properties.title = "Especificação de Requisitos"
    doc.core_properties.subject = "VinceArt"
    doc.core_properties.comments = "Gerado a partir de Requisitos/URS.md"

    headings = [payload for kind, payload in blocks if kind == "heading" and 1 < payload[0] <= 3]

    meta = {}
    capa_feita = False
    i = 0
    while i < len(blocks):
        kind, payload = blocks[i]
        i += 1

        if kind == "heading":
            level, text = payload
            if level == 1:
                meta["__titulo__"] = text
                continue
            m = RF_TITLE_RE.match(text) if level >= 4 else None
            if m and i < len(blocks) and blocks[i][0] == "list":
                render_requisito(doc, m.group(1), m.group(2), blocks[i][1])
                i += 1
                continue
            p = doc.add_heading(level=min(level - 1, 9))
            add_inline(p, text)

        elif kind == "para":
            campos = None if capa_feita else parse_meta(payload)
            if campos:
                meta.update(campos)
                render_capa(doc, meta, logo_path)
                add_toc(doc, headings)
                capa_feita = True
            elif payload == ["&nbsp;"]:
                doc.add_paragraph()          # espaçador explícito, para assinaturas
            else:
                p = doc.add_paragraph()
                add_inline(p, " ".join(payload))

        elif kind == "list":
            render_list(doc, payload)
        elif kind == "table":
            render_table(doc, *payload)
        elif kind == "code":
            render_code(doc, payload)

    add_footer(doc)
    doc.save(out_path)


def main():
    src, dst = sys.argv[1], sys.argv[2]
    with open(src, encoding="utf-8") as fh:
        lines = fh.readlines()
    logo = os.path.join(os.path.dirname(os.path.abspath(src)), "logo-biopark.png")
    render(parse_blocks(lines), dst, logo)
    print(f"gerado: {dst}")


if __name__ == "__main__":
    main()
